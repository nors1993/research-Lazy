"""DOCX document generator for patent and paper writing."""

from dataclasses import dataclass
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx import Document

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .constants import FONTS, PAGE_SIZES, PAPER_MARGINS, PATENT_MARGINS


@dataclass
class DocxConfig:
    """Configuration for DOCX generation."""

    page_width: float = 21.0
    page_height: float = 29.7
    top_margin: float = 2.54
    bottom_margin: float = 2.54
    left_margin: float = 2.54
    right_margin: float = 2.54

    body_font: str = "Times New Roman"
    heading_font: str = "Arial"
    chinese_body_font: str = "SimSun"
    chinese_heading_font: str = "SimHei"

    body_font_size: int = 12
    heading1_size: int = 22
    heading2_size: int = 16
    heading3_size: int = 14

    doc_type: str = "PAPER"
    language: str = "zh-CN"

    @classmethod
    def for_patent(cls, language: str = "zh-CN") -> "DocxConfig":
        config = cls(
            page_width=PAGE_SIZES["A4"]["width"],
            page_height=PAGE_SIZES["A4"]["height"],
            top_margin=PATENT_MARGINS["top"],
            bottom_margin=PATENT_MARGINS["bottom"],
            left_margin=PATENT_MARGINS["left"],
            right_margin=PATENT_MARGINS["right"],
            body_font=FONTS["chinese"]["body"],
            heading_font=FONTS["chinese"]["heading"],
            body_font_size=12,
            heading1_size=22,
            heading2_size=16,
            heading3_size=14,
            doc_type="PATENT",
            language=language,
        )
        return config

    @classmethod
    def for_paper(cls, language: str = "en-US") -> "DocxConfig":
        config = cls(
            page_width=PAGE_SIZES["A4"]["width"],
            page_height=PAGE_SIZES["A4"]["height"],
            top_margin=PAPER_MARGINS["top"],
            bottom_margin=PAPER_MARGINS["bottom"],
            left_margin=PAPER_MARGINS["left"],
            right_margin=PAPER_MARGINS["right"],
            body_font=FONTS["english"]["body"],
            heading_font=FONTS["english"]["heading"],
            body_font_size=12,
            heading1_size=16,
            heading2_size=14,
            heading3_size=12,
            doc_type="PAPER",
            language=language,
        )
        return config


class DocxGenerator:
    """Generator for creating professional DOCX documents."""

    def __init__(self, config: DocxConfig | None = None):
        self.config = config or DocxConfig.for_paper()

    def create_document(self):
        """Create a new DOCX document with proper formatting."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(self.config.page_width)
        section.page_height = Cm(self.config.page_height)
        section.top_margin = Cm(self.config.top_margin)
        section.bottom_margin = Cm(self.config.bottom_margin)
        section.left_margin = Cm(self.config.left_margin)
        section.right_margin = Cm(self.config.right_margin)
        return doc

    def add_heading(self, doc, text: str, level: int = 1) -> None:
        """Add a heading with proper formatting."""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = self.config.heading_font
            if self.config.language == "zh-CN":
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.config.heading_font)

            if level == 1:
                run.font.size = Pt(self.config.heading1_size)
                run.font.bold = True
            elif level == 2:
                run.font.size = Pt(self.config.heading2_size)
                run.font.bold = True
            else:
                run.font.size = Pt(self.config.heading3_size)
                run.font.bold = True

        if level == 1 and self.config.doc_type == "PATENT":
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_paragraph(self, doc, text: str, indent: bool = False) -> None:
        """Add a paragraph with proper formatting."""
        para = doc.add_paragraph()
        if indent:
            para.paragraph_format.first_line_indent = Cm(0.74)

        run = para.add_run(text)
        run.font.name = self.config.body_font
        if self.config.language == "zh-CN":
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.config.chinese_body_font)
        run.font.size = Pt(self.config.body_font_size)
        para.paragraph_format.line_spacing = Pt(22)
        para.paragraph_format.space_after = Pt(0)

    def add_claim(self, doc, claim_number: int, text: str, is_independent: bool = False) -> None:
        """Add a patent claim with proper formatting."""
        para = doc.add_paragraph()
        num_run = para.add_run(f"{claim_number}. ")
        num_run.font.name = self.config.body_font
        num_run.font.bold = True
        num_run.font.size = Pt(self.config.body_font_size)

        content_run = para.add_run(text)
        content_run.font.name = self.config.body_font
        if self.config.language == "zh-CN":
            content_run._element.rPr.rFonts.set(qn("w:eastAsia"), self.config.chinese_body_font)
        content_run.font.size = Pt(self.config.body_font_size)
        para.paragraph_format.line_spacing = Pt(22)
        para.paragraph_format.space_after = Pt(6)

    def add_table(self, doc, data: list[list[str]], headers: list[str] | None = None) -> None:
        """Add a table with proper formatting."""
        if not data:
            return

        col_count = len(headers) if headers else len(data[0]) if data else 0
        table = doc.add_table(rows=0, cols=col_count)
        table.style = "Table Grid"

        if headers:
            header_row = table.add_row()
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.name = self.config.body_font

        for row_data in data:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                row.cells[i].text = str(cell_text)

    def set_document_title(self, doc, title: str) -> None:
        """Set the document title."""
        title_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title_para.add_run(title)
        run.font.name = self.config.heading_font
        if self.config.language == "zh-CN":
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.config.chinese_heading_font)
        run.font.size = Pt(self.config.heading1_size)
        run.font.bold = True
        title_para.paragraph_format.space_after = Pt(24)


class PatentDocxGenerator(DocxGenerator):
    """Specialized generator for Chinese patent documents."""

    def __init__(self):
        super().__init__(DocxConfig.for_patent())

    def create_patent(
        self,
        title: str,
        abstract: str,
        claims: list[dict[str, Any]],
        technical_field: str,
        background: str,
        invention_content: str,
        detailed_description: str,
        beneficial_effects: str = "",
        output_path: str | None = None,
    ) -> str:
        """Create a complete patent document."""
        doc = self.create_document()

        self.set_document_title(doc, title)
        self.add_heading(doc, "说明书摘要", level=1)
        self.add_paragraph(doc, abstract, indent=True)

        self.add_heading(doc, "权利要求书", level=1)
        for i, claim in enumerate(claims, 1):
            is_independent = claim.get("independent", False)
            self.add_claim(doc, i, claim["text"], is_independent)

        self.add_heading(doc, "技术领域", level=1)
        self.add_paragraph(doc, technical_field, indent=True)

        self.add_heading(doc, "背景技术", level=1)
        self.add_paragraph(doc, background, indent=True)

        self.add_heading(doc, "发明内容", level=1)
        self.add_paragraph(doc, invention_content, indent=True)

        if beneficial_effects:
            self.add_heading(doc, "有益效果", level=2)
            self.add_paragraph(doc, beneficial_effects, indent=True)

        self.add_heading(doc, "具体实施方式", level=1)
        self.add_paragraph(doc, detailed_description, indent=True)

        if output_path:
            doc.save(output_path)

        return output_path or ""


class PaperDocxGenerator(DocxGenerator):
    """Specialized generator for academic papers."""

    def __init__(self, language: str = "en-US"):
        config = DocxConfig.for_paper(language)
        super().__init__(config)

    def create_paper(
        self,
        title: str,
        abstract: str,
        keywords: list[str],
        sections: dict[str, str],
        references: list[dict] | None = None,
        output_path: str | None = None,
    ) -> str:
        """Create a complete academic paper."""
        doc = self.create_document()

        self.set_document_title(doc, title)

        abstract_label = "Abstract" if self.config.language != "zh-CN" else "摘要"
        self.add_heading(doc, abstract_label, level=1)
        self.add_paragraph(doc, abstract, indent=True)

        if keywords:
            kw_label = "Keywords" if self.config.language != "zh-CN" else "关键词"
            flat_keywords = []
            for kw in keywords:
                if isinstance(kw, list):
                    flat_keywords.extend(kw)
                else:
                    flat_keywords.append(str(kw))
            self.add_paragraph(doc, f"{kw_label}: {', '.join(flat_keywords)}")

        section_titles = {
            "introduction": "1. Introduction" if self.config.language != "zh-CN" else "1 引言",
            "related_work": "2. Related Work" if self.config.language != "zh-CN" else "2 相关工作",
            "methodology": "3. Methodology" if self.config.language != "zh-CN" else "3 方法",
            "results": "4. Results" if self.config.language != "zh-CN" else "4 结果",
            "discussion": "5. Discussion" if self.config.language != "zh-CN" else "5 讨论",
            "conclusion": "6. Conclusion" if self.config.language != "zh-CN" else "6 结论",
        }

        for section_key, section_content in sections.items():
            section_title = section_titles.get(section_key, section_key.upper())
            self.add_heading(doc, section_title, level=2)

            if isinstance(section_content, str):
                paragraphs = section_content.split("\n\n")
                for para in paragraphs:
                    if para.strip():
                        self.add_paragraph(doc, para.strip(), indent=True)
            else:
                self.add_paragraph(doc, str(section_content), indent=True)

        if references:
            ref_label = "References" if self.config.language != "zh-CN" else "参考文献"
            self.add_heading(doc, ref_label, level=1)
            for i, ref in enumerate(references, 1):
                ref_text = self._format_reference(ref)
                self.add_paragraph(doc, f"[{i}] {ref_text}")

        if output_path:
            doc.save(output_path)

        return output_path or ""

    def _format_reference(self, ref: dict) -> str:
        """Format a reference entry."""
        authors = ref.get("authors", [])
        if isinstance(authors, list):
            flat_authors = []
            for a in authors:
                if isinstance(a, list):
                    flat_authors.extend(a)
                else:
                    flat_authors.append(str(a))
            author_str = ", ".join(flat_authors[:3])
            if len(flat_authors) > 3:
                author_str += ", et al."
        else:
            author_str = str(authors)

        year = ref.get("year", "")
        title = ref.get("title", "")
        journal = ref.get("journal", ref.get("venue", ""))
        volume = ref.get("volume", "")
        pages = ref.get("pages", "")

        parts = [author_str, title]
        if journal:
            parts.append(journal)
        if volume:
            parts.append(f"({volume})")
        if year:
            parts.append(str(year))
        if pages:
            parts.append(f"pp. {pages}")

        return ". ".join(parts)
