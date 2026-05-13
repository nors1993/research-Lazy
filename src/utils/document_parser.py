"""Document parser for converting .docx, .pdf, .md files to markdown format."""

import re
from pathlib import Path

from docx import Document

from .logger import get_logger

logger = get_logger(__name__)


class DocumentParser:
    """Parse document files to markdown format."""

    SUPPORTED_FORMATS = {".docx", ".pdf", ".md", ".txt", ".doc"}

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file format is supported."""
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_FORMATS

    @classmethod
    def parse(cls, file_path: str) -> str:
        """Parse document file to markdown format.

        Args:
            file_path: Path to the document file

        Returns:
            Markdown content string

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        if ext == ".docx":
            return cls._parse_docx(path)
        elif ext == ".pdf":
            return cls._parse_pdf(path)
        elif ext == ".doc":
            return cls._parse_doc(path)
        elif ext in (".md", ".txt"):
            return cls._parse_markdown(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Parse .docx file to markdown."""
        try:
            doc = Document(str(path))
            markdown_lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect heading level based on style
                style_name = para.style.name.lower() if para.style and para.style.name else ""

                if "heading 1" in style_name:
                    markdown_lines.append(f"# {text}")
                elif "heading 2" in style_name:
                    markdown_lines.append(f"## {text}")
                elif "heading 3" in style_name:
                    markdown_lines.append(f"### {text}")
                elif "heading 4" in style_name:
                    markdown_lines.append(f"#### {text}")
                elif "heading 5" in style_name:
                    markdown_lines.append(f"##### {text}")
                elif "heading 6" in style_name:
                    markdown_lines.append(f"###### {text}")
                else:
                    # Check for bold/italic formatting in runs
                    formatted_text = text
                    for run in para.runs:
                        if run.bold and run.text:
                            formatted_text = formatted_text.replace(
                                run.text, f"**{run.text}**"
                            )
                        if run.italic and run.text:
                            formatted_text = formatted_text.replace(
                                run.text, f"*{run.text}*"
                            )

                    # Clean up any double formatting
                    formatted_text = formatted_text.replace("****", "")
                    formatted_text = formatted_text.replace("****", "")

                    markdown_lines.append(formatted_text)

            # Process tables
            for table in doc.tables:
                markdown_lines.append("")
                markdown_lines.append("| " + " | ".join([""] * len(table.columns)) + " |")
                markdown_lines.append("| " + " | ".join(["---"] * len(table.columns)) + " |")

                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace("\n", " ")
                        cells.append(cell_text)
                    markdown_lines.append("| " + " | ".join(cells) + " |")

                markdown_lines.append("")

            return "\n\n".join(markdown_lines)

        except Exception as e:
            logger.error("docx_parse_failed", path=str(path), error=str(e))
            raise

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Parse .pdf file to markdown using pdfplumber."""
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                markdown_lines = []

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        # Clean up text
                        text = text.strip()
                        # Remove excessive whitespace
                        text = re.sub(r"\n{3,}", "\n\n", text)
                        # Add page separator for multi-page documents
                        if page_num > 1:
                            markdown_lines.append(f"\n---\n**Page {page_num}**\n")
                        markdown_lines.append(text)

                result = "\n\n".join(markdown_lines)

                if not result.strip():
                    logger.warning("pdf_no_text_extracted", path=str(path))
                    return "<!-- PDF contains no extractable text -->"

                return result

        except ImportError:
            logger.error("pdfplumber_not_installed")
            raise ImportError(
                "pdfplumber is required for PDF parsing. Install with: pip install pdfplumber"
            )
        except Exception as e:
            logger.error("pdf_parse_failed", path=str(path), error=str(e))
            raise

    @staticmethod
    def _parse_markdown(path: Path) -> str:
        """Parse .md or .txt file - return content as-is."""
        try:
            with open(path, encoding="utf-8") as f:
                content = f.read()
            return content.strip()
        except Exception as e:
            logger.error("markdown_parse_failed", path=str(path), error=str(e))
            raise

    @staticmethod
    def _parse_doc(path: Path) -> str:
        """Parse .doc file to markdown using textract."""
        try:
            import textract

            text = textract.process(str(path))
            # textract returns bytes, decode to string
            content = text.decode("utf-8", errors="ignore")
            return content.strip()
        except ImportError:
            logger.error("textract_not_installed")
            raise ImportError(
                "textract is required for .doc parsing. Install with: pip install textract"
            )
        except Exception as e:
            logger.error("doc_parse_failed", path=str(path), error=str(e))
            raise

    @classmethod
    def parse_from_bytes(cls, content: bytes, filename: str, original_ext: str) -> str:
        """Parse document from bytes content.

        Args:
            content: File content as bytes
            filename: Original filename
            original_ext: Original file extension (e.g., ".docx")

        Returns:
            Markdown content string
        """
        import tempfile

        # Create temp file with correct extension
        with tempfile.NamedTemporaryFile(
            suffix=original_ext, delete=False
        ) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            return cls.parse(tmp_path)
        finally:
            # Clean up temp file
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass
